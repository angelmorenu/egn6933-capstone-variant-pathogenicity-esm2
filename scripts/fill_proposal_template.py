from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def _insert_paragraph_after(paragraph, text: str, style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    new_para.add_run(text)
    return new_para


def _replace_checkbox(text: str) -> str:
    return text.replace("[ ]", "[x]", 1) if "[ ]" in text else text


def main() -> None:
    template = Path("/Users/angelhdmorenu/Downloads/Project Proposal Template.docx")
    out = Path(
        "/Users/angelhdmorenu/Desktop/EGN 6933 – Project in Applied Data Science/"
        "Machine Learning Classification of Pathogenic vs. Benign Missense Variants Using Protein Language Model Embeddings/"
        "project-proposal/Morenu_Project_Proposal_Template_Filled_v2.docx"
    )

    ans = {
        "project_name": "Machine Learning Classification of Pathogenic vs. Benign Missense Variants Using Protein Language Model Embeddings",
        "team_lead": "Angel Morenu",
        "team_members": "N/A (Individual Project)",
        "problem": (
            "Develop an end-to-end, reproducible ML pipeline that classifies missense variants as pathogenic vs. benign using fixed-length "
            "ESM2 embedding vectors as features. The emphasis is on prediction quality and leakage-aware evaluation (gene/protein-disjoint splits), "
            "not on mechanistic interpretation."
        ),
        "stakeholders": (
            "Rare disease researchers and clinical genomics teams who prioritize variants for follow-up; computational biology groups who need "
            "scalable, reproducible variant-scoring workflows."
        ),
        "ethical": (
            "Uses de-identified ClinVar variant records and a lab-provided, post-QC ClinVar-derived missense-only table with aligned embeddings. "
            "No private patient information is used. The model is a research tool for prioritization and explicitly not for standalone clinical diagnosis; "
            "labeling follows ClinVar clinical significance conventions and excludes ambiguous/conflicting categories."
        ),
        "data_source": (
            "ClinVar (https://www.ncbi.nlm.nih.gov/clinvar/) for authoritative provenance/mapping, plus a Dr. Fan lab-provided, post-QC missense-only "
            "ClinVar-derived table (Dylan Tan) and aligned precomputed ESM2 embeddings."
        ),
        "data_access": "Yes (ClinVar is public; the derived table/embeddings are available through Dr. Fan’s lab).",
        "data_description": (
            "A labeled missense-only variant table joined to high-dimensional embedding features (ESM2 fixed-length vectors). A validated 5,000-row pilot "
            "artifact is already produced to confirm end-to-end alignment; the full dataset scale depends on label filtering and curation."
        ),
        "methods": (
            "Representation learning via pretrained protein language model embeddings (ESM2) as input features; baseline classifiers (Logistic Regression, "
            "Random Forest) and an optional shallow MLP; leakage-aware evaluation with group-aware splitting by gene/protein identifiers and calibrated "
            "probability outputs."
        ),
        "pipeline": (
            "Scripted data build steps: apply strict label policy (P/LP vs B/LB; exclude VUS/conflicting), retain missense-only variants (from post-QC table "
            "and/or QC checks), align each variant row to its embedding vector, and write versioned artifacts (Parquet/TSV + NumPy embeddings + metadata + "
            "saved split indices)."
        ),
        "repro": (
            "GitHub version control; pinned Python environment (Conda/venv); deterministic scripts with fixed seeds; artifact metadata (ClinVar release identifiers, "
            "filtering choices, embedding dimension/model, split seeds/group keys) to enable repeatable training and evaluation."
        ),
        "interface": "Streamlit dashboard for interactive scoring + a CLI for batch scoring of variant tables (CSV/VCF-derived exports).",
        "core_func": "Input a missense variant (single or batch) and receive a calibrated pathogenicity probability with a ranked output table for prioritization.",
        "metrics": (
            "Primary: AUROC and AUPRC (robust under class imbalance). Secondary: precision/recall/F1/balanced accuracy; calibration checks (Brier score and "
            "calibration curves)."
        ),
        "tests": (
            "DeLong test for AUROC comparisons; paired bootstrap/permutation tests for AUPRC; bootstrapped confidence intervals on a held-out gene/protein-disjoint "
            "test set."
        ),
        "w1_4": "Data acquisition + strict label policy; build reproducible dataset artifacts and sanity checks; finalize gene/protein-aware splitting plan and curated dataset artifact.",
        "w5_8": "Train baselines (Logistic Regression, Random Forest); hyperparameter tuning; threshold selection; initial AUROC/AUPRC and calibration assessment.",
        "w9_12": "Optional shallow MLP; rigorous statistical comparisons (paired tests + CIs); error analysis and model calibration refinement.",
        "w13_15": "Deploy Streamlit + CLI scoring; package model + evaluation report; finalize documentation and present demo.",
        "new_knowledge": "Protein language model embeddings (ESM2) applied to human variant classification, including leakage-aware gene/protein-disjoint evaluation and probability calibration.",
    }

    doc = Document(str(template))

    for p in list(doc.paragraphs):
        t = p.text.strip()

        # Remove emoji that can break LaTeX-based PDF conversion
        if "📋" in p.text:
            p.text = p.text.replace("📋 ", "").replace("📋", "")

        if t.startswith("Project Name:"):
            p.text = f"Project Name: {ans['project_name']}"
            continue
        if t.startswith("Team Lead:"):
            p.text = f"Team Lead: {ans['team_lead']}"
            continue
        if t.startswith("Team Members:"):
            p.text = f"Team Members: {ans['team_members']}"
            continue

        if t.startswith("The Problem:"):
            _insert_paragraph_after(p, ans["problem"])
            continue
        if t.startswith("Stakeholders:"):
            _insert_paragraph_after(p, ans["stakeholders"])
            continue
        if t.startswith("Societal/Ethical Context:"):
            _insert_paragraph_after(p, ans["ethical"])
            continue

        if t.startswith("Data Source:"):
            _insert_paragraph_after(p, ans["data_source"])
            continue
        if t.startswith("Status:"):
            _insert_paragraph_after(p, ans["data_access"])
            continue
        if t.startswith("Description:"):
            _insert_paragraph_after(p, ans["data_description"])
            continue

        if t.startswith("Methods/Algorithms:"):
            _insert_paragraph_after(p, ans["methods"])
            continue
        if t.startswith("Data Pipeline:"):
            _insert_paragraph_after(p, ans["pipeline"])
            continue
        if t.startswith("Reproducibility:"):
            _insert_paragraph_after(p, ans["repro"])
            continue

        if t.startswith("The Interface:"):
            _insert_paragraph_after(p, ans["interface"])
            continue
        if t.startswith("Core Functionality:"):
            _insert_paragraph_after(p, ans["core_func"])
            continue

        if t.startswith("Evaluation Metrics:"):
            _insert_paragraph_after(p, ans["metrics"])
            continue
        if t.startswith("Statistical Tests:"):
            _insert_paragraph_after(p, ans["tests"])
            continue

        if t.startswith("Weeks 1–4:"):
            _insert_paragraph_after(p, ans["w1_4"])
            continue
        if t.startswith("Weeks 5–8:"):
            _insert_paragraph_after(p, ans["w5_8"])
            continue
        if t.startswith("Weeks 9–12:"):
            _insert_paragraph_after(p, ans["w9_12"])
            continue
        if t.startswith("Weeks 13–15:"):
            _insert_paragraph_after(p, ans["w13_15"])
            continue

        if t.startswith('What is one major skill, library,'):
            _insert_paragraph_after(p, ans["new_knowledge"])
            continue

        if t.startswith("[ ]"):
            p.text = _replace_checkbox(p.text)
            continue

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
